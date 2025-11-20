# insights_agent.py
# DigitalABCs Compliant AI Agent — Gemini API Version
# Empowerment | Compliance | Practicality | Privacy

# -----------------------------------------------------------------
# --- 🟢 START: ENVIRONMENT FIX FOR CPANEL 🟢 ---
# -----------------------------------------------------------------
# These lines MUST be at the very top, before any other imports.
# They fix the "pthread_create failed" and "can't start new thread" errors
# by telling Google's gRPC library to use a simpler, thread-safe method.
import os
os.environ["GRPC_POLL_STRATEGY"] = "poll"
os.environ["GRPC_ENABLE_FORK_SUPPORT"] = "0"
# -----------------------------------------------------------------
# --- 🟢 END: ENVIRONMENT FIX 🟢 ---
# -----------------------------------------------------------------


# -----------------------------------------------------------------
# --- IMPORTS ---
# -----------------------------------------------------------------
import logging
from google import genai
import sys
from pathlib import Path

# --- ▼▼▼ START: SERVICE ACCOUNT AUTHENTICATION ▼▼▼ ---
# This block MUST come before you import 'vertexai'
SCRIPT_DIR = Path(__file__).resolve().parent

# This should be the name of the JSON key file you uploaded
KEY_FILE_PATH = SCRIPT_DIR / "gen-lang-client-0944750405-a5911516ca25.json"

if not KEY_FILE_PATH.exists():
    print(f"❌ Error: Authentication key file not found at {KEY_FILE_PATH}")
    print("❌ Please download your Google Cloud service account JSON key,")
    print("❌ upload it to this directory, and ensure the filename matches.")
    sys.exit(1)
    
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = str(KEY_FILE_PATH)
# --- ▲▲▲ END: SERVICE ACCOUNT AUTHENTICATION ▲▲▲ ---


import json
from datetime import datetime, timedelta
from typing import List, Annotated
import shutil  # To copy the final HTML file
import base64  # To create the placeholder image on failure
from html import escape # To sanitize HTML output

# ADDED: Vertex AI Imports for Image Generation
import vertexai
try:
    from vertexai.preview.vision_models import ImageGenerationModel
except ImportError:
    print("WARNING: vertexai library not fully found. Image generation might fail.")
    ImageGenerationModel = None

# FIX: Attempt to import orjson for robust JSON decoding
try:
    import orjson
    JSON_LOADER = orjson.loads
except ImportError:
    JSON_LOADER = json.loads


# Load environment variables
from dotenv import load_dotenv
load_dotenv()

# Structured output for trust & simplicity
from pydantic import BaseModel, Field

# Gemini SDK
import google.generativeai as genai
from google.generativeai import types 

# Pure Chroma for private RAG (no LangChain)
import chromadb
from chromadb.utils import embedding_functions

# Document & news handling
from dateutil import parser as date_parser
import feedparser
from docx import Document as DocxDocument

# Practical outputs
import pandas as pd

# Accessibility audit (Optional, can be uncommented in main)
# from selenium import webdriver
# from axe_selenium_python import Axe

# Ensure clean environment
os.environ["TOKENIZERS_PARALLELISM"] = "false"

# -----------------------------------------------------------------
# --- PATH & FILENAME CONFIGURATION ---
# -----------------------------------------------------------------
# (SCRIPT_DIR is already defined at the top)

# Paths for local files (brand voice, vector DB)
BRAND_VOICE_FILE = SCRIPT_DIR / "brand_voice.docx"
CHROMA_DB_PATH = SCRIPT_DIR / "chroma_db"

# Path for your private ARCHIVES (docx, xlsx, and a copy of the html)
ARCHIVE_OUTPUT_DIR = Path("/home/digitala/my_scripts/agents/my_business_insights/outputs")
ARCHIVE_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Path for your PUBLIC website
PUBLIC_HTML_DIR = Path("/home/digitala/public_html/insights")
PUBLIC_HTML_DIR.mkdir(parents=True, exist_ok=True)

# Generate timestamp like: 2025-10-21_14-30
timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")

# -----------------------------------------------------------------
# --- NEWS SOURCES ---
# -----------------------------------------------------------------
NEWS_SOURCES = [
    "https://www.smartcompany.com.au/feed/",
    "https://businessnewsaustralia.com/feed/",
    "https://asic.gov.au/about-asic/media-centre/news-releases/media-releases-2025/feed/",
    "https://ausbiz.com.au/feeds/rss",
    "https://www.itnews.com.au/rss",
    "https://www.crn.com.au/rss",
    "https://www.themandarin.com.au/feed/",
    "https://www.businessnewsaustralia.com/articles/rss/startup-daily",
    "https://www.business.gov.au/news.rss",
    "https://www.dynamicbusiness.com.au/feed",
    "https://www.businessnewsaustralia.com/articles/rss/sme",
    "https://www.asbfeo.gov.au/news/rss.xml"
]

# -----------------------------------------------------------------
# --- API & CLIENT CONFIGURATION ---
# -----------------------------------------------------------------


# Get Gemini API key from environment
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    print("❌ Error: GEMINI_API_KEY not found in .env file")
    sys.exit(1)
os.environ["GOOGLE_API_KEY"] = GEMINI_API_KEY
# Configure Gemini to use REST (thread-safe for cPanel)
genai.configure(
    api_key=GEMINI_API_KEY,
    transport="rest"
)

# Create the Google embedding function for ChromaDB
try:
    # 🟢 NOTE: We can't pass 'transport="rest"' here,
    # so the os.environ variables at the top of the file
    # are CRITICAL for this to work.
    google_ef = embedding_functions.GoogleGenerativeAiEmbeddingFunction(
        api_key=GEMINI_API_KEY,
        model_name="models/text-embedding-004",
        task_type="RETRIEVAL_DOCUMENT"
    )
except Exception as e:
    print(f"❌ FAILED to create embedding function: {e}")
    print("❌ This often means the 'Generative Language API' is not enabled in your Google Cloud project.")
    sys.exit(1)


# === 🟢 NOTE: Vertex AI is NO LONGER initialized here. ===
# It is now initialized locally inside generate_ai_image() to prevent auth conflicts.


# === STEP 1: LOAD BRAND VOICE ===
def load_brand_voice() -> str:
    brand_file = BRAND_VOICE_FILE
    if not brand_file.exists():
        print(f"❌ Error: '{BRAND_VOICE_FILE}' not found.")
        sys.exit(1)
    try:
        doc = DocxDocument(brand_file)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()]) or "You are a trusted Australian small business advisor."
    except Exception as e:
        print(f"❌ Error reading brand_voice.docx: {e}")
        sys.exit(1)

# === STEP 2: PRIVATE KNOWLEDGE BASE (PURE CHROMA) ===
def create_compliance_kb():
    """Create a local, private knowledge base using ChromaDB."""
    client = chromadb.PersistentClient(path=str(CHROMA_DB_PATH))
    
    collection = client.get_or_create_collection(
        name="compliance_kb",
        embedding_function=google_ef, # Uses the corrected function
        metadata={"hnsw:space": "cosine"}
    )
    
    # Build knowledge base (only once)
    if collection.count() == 0:
        print("🔄 Building new knowledge base... (This may take a moment for API embeddings)")
        brand_text = load_brand_voice()
        compliance_text = """
        Australian Privacy Principle 1: Open and transparent management of personal information.
        Disability Discrimination Act 1992: Digital services must be accessible.
        Small businesses must protect personal information from misuse.
        Always use plain Australian English. Avoid jargon. Focus on actionable steps.
        """
        all_text = brand_text + "\n\n" + compliance_text
        
        # Simple chunking
        chunks = [all_text[i:i+500] for i in range(0, len(all_text), 450)]
        
        try:
            collection.add(
                documents=chunks,
                ids=[f"chunk_{i}" for i in range(len(chunks))]
            )
            print("✅ Knowledge base created successfully via API.")
        except Exception as e:
            print(f"❌ API Error during collection.add(): {e}")
            print("❌ Please check your GEMINI_API_KEY and ensure the 'Generative Language API' is enabled in your Google Cloud project.")
            sys.exit(1)
    else:
        print("✅ Knowledge base loaded from disk.")
    
    return collection

def retrieve_context(collection, query: str, k: int = 2) -> str:
    """Retrieve relevant brand/compliance context."""
    results = collection.query(query_texts=[query], n_results=k)
    return "\n".join(results['documents'][0])

# === STEP 3: FETCH AUSTRALIAN NEWS (UPDATED FOR 7 DAYS) ===
def fetch_recent_news():
    cutoff = datetime.now() - timedelta(days=7) 
    articles = []
    for url in NEWS_SOURCES:
        try:
            feed = feedparser.parse(url)
            for entry in feed.entries[:4]:
                try:
                    pub_date_str = getattr(entry, 'published', None) or getattr(entry, 'updated', None)
                    if not pub_date_str:
                        continue
                    
                    pub_date = date_parser.parse(pub_date_str)
                    if pub_date.tzinfo:
                        pub_date = pub_date.replace(tzinfo=None)
                    if pub_date >= cutoff:
                        articles.append({
                            "title": entry.title,
                            "summary": entry.summary,
                            "link": entry.link,
                            "date": pub_date.strftime("%Y-%m-%d")
                        })
                except:
                    continue
        except Exception as e:
            print(f"⚠️  Warning: {url} - {e}")
            continue
    return articles[:6]

# === (REVISED) STEP 3.5: TEXT-TO-IMAGE FUNCTION ===
def generate_ai_image(prompt: str, image_filename: str) -> bool:
    """
    Generates an image using Vertex AI (Imagen) and saves it to the public insights folder.
    Returns True on success, False on failure.
    """
    print(f"🖼️  Attempting to generate AI image for prompt: '{prompt}'")
    
    if ImageGenerationModel is None:
        logging.error("❌ 'vertexai' library is missing or failed to import.")
        return False

    try:
        # 1. 🟢 AUTHENTICATE & INITIALIZE (MOVED HERE) 🟢
        # This is now local to this function to prevent auth conflicts
        GCP_PROJECT_ID = os.getenv("GOOGLE_PROJECT_ID")
        GCP_REGION = os.getenv("GOOGLE_LOCATION", "us-central1")

        if not GCP_PROJECT_ID or not GCP_REGION:
            logging.error("❌ GOOGLE_PROJECT_ID or GOOGLE_LOCATION not set in .env")
            return False

        vertexai.init(project=GCP_PROJECT_ID, location=GCP_REGION)
        print(f"✅ Vertex AI initialized locally for project: {GCP_PROJECT_ID}")

        # 2. CREATE THE MODEL AND GENERATE
        model = ImageGenerationModel.from_pretrained("imagegeneration@006")
        
        response = model.generate_images(
            prompt=prompt,
            number_of_images=1,
            aspect_ratio="16:9",
        )

        # 3. SAVE THE IMAGE
        if response and response.images:
            image_data = base64.b64decode(response.images[0]._image_bytes)
            
            save_path = PUBLIC_HTML_DIR / image_filename 
            
            with open(save_path, "wb") as f:
                f.write(image_data)
            
            print(f"✅ Image saved successfully: {save_path}")
            return True
        else:
            logging.warning("⚠️ Vertex AI returned no images.")
            return False

    except Exception as e:
        # This will catch any auth errors, API errors, etc.
        logging.error(f"❌ Vertex AI Image Generation FAILED: {e}")
        return False
# === END STEP 3.5 ===


# === (FIXED) STEP 4: STRUCTURED OUTPUT MODEL ===
# 🟢 FIX: Added 'tile_summary' and 'min_length' constraints
class BusinessInsight(BaseModel):
    title: str = Field(
        ..., 
        min_length=10,
        description="e.g., 'Australian Small Business Weekly: 21 Oct 2025'"
    )
    
    tile_summary: str = Field(
        ..., 
        min_length=20,
        description="A very short, 1-2 sentence hook for the insights grid. Max 160 characters. Must be compelling and make the user want to click."
    )
    
    summary: str = Field(
        ..., 
        min_length=50,
        description="3-4 paragraphs (no bullet points) focusing on THIS WEEK's specific developments, framed around AI, automation, and immediate opportunities for micro-businesses."
    )
    key_updates: Annotated[List[str], Field(
        min_length=3, max_length=4, 
        description="Specific actionable, time-sensitive tips for micro-businesses, written in plain language, emphasizing AI, automation, or software agents based on THIS WEEK's news."
    )]
    image_generation_prompt: str = Field(
        ..., 
        min_length=20,
        description="A 1-2 sentence prompt for an AI image model (like Imagen) to generate a relevant 16:9 image. Must be descriptive."
    )
    image_alt_text: str = Field(
        ..., 
        min_length=10,
        description="WCAG-compliant alt text for the generated image. Must be descriptive and context-rich."
    )
    sources: Annotated[List[str], Field(
        min_length=2, 
        description="Source URLs"
    )]
# === END STEP 4 ===

# === (FIXED) STEP 5: GENERATE INSIGHTS (NATIVE GEMINI SDK) ===
def generate_insights(news_articles, kb_collection):
    """
    Calls Gemini to get structured text, then calls Imagen to generate an image.
    🟢 FIX: Returns (insight, hook_text, image_filename)
    """
    
    # 🟢 NOTE: This is the call that was failing.
    # It runs first, using the API Key (via 'google_ef').
    print("🔄 Retrieving brand context from ChromaDB...")
    context = retrieve_context(kb_collection, "Australian small business compliance and brand voice")
    
    print("🔄 Calling Gemini API (text)...")
    
    # Prepare news
    news_text = "\n\n".join([
        f"Title: {a['title']}\nDate: {a['date']}\nSummary: {a['summary']}\nSource: {a['link']}"
        for a in news_articles
    ]) if news_articles else "No recent updates."
    
    # Get Pydantic model schema for JSON output
    schema_data = BusinessInsight.model_json_schema()
    
    # 🟢 FIX: Added 'tile_summary' and fixed the closing brace
    json_schema = {
        "type": "object",
        "properties": {
            "title": schema_data["properties"]["title"],
            "tile_summary": schema_data["properties"]["tile_summary"], 
            "summary": schema_data["properties"]["summary"],
            "key_updates": schema_data["properties"]["key_updates"],
            "image_generation_prompt": schema_data["properties"]["image_generation_prompt"],
            "image_alt_text": schema_data["properties"]["image_alt_text"],
            "sources": schema_data["properties"]["sources"]
        }, # <--- This brace was missing
        "required": [
            "title", "tile_summary", "summary", "key_updates", 
            "image_generation_prompt", "image_alt_text", "sources"
        ]
    }
    
    # Define the content of the emotional hook/CTA
    hook_and_cta_text = (
        "Feeling overwhelmed? You're not alone. I built this business while managing Bipolar and ADHD, "
        "turning chaos into clarity. These key updates aren't just news; they're opportunities. "
        "Let DigitalABCs show you how to use 'Asynchronous Automation' to build a business that works *with* your brain, not against it."
    )

    prompt = f"""[Your Guidelines]
{context}

[This Week's News - Last 7 Days]
{news_text}

You are a trusted advisor for **DigitalABCs** to Australian micro-businesses (<5 employees). Your voice is relatable, supportive, and empowering (based on Trish's story of resilience).
Generate a weekly insight that is practical, empowering, and grounded in Australian regulations.
**Crucially, frame all insights and key takeaways through the lens of AI, automation, and the use of simple software agents to simplify life and give power back to employees.**
Focus ONLY on developments from the past 7 days. Avoid jargon.

IMPORTANT: Respond with ONLY valid JSON that matches this exact structure:
{json.dumps(json_schema, indent=2)}

DO NOT include any text before or after the JSON. DO NOT use markdown code blocks. Just output the raw JSON."""

    # Configure Gemini Model
    model = genai.GenerativeModel(
        'models/gemini-2.5-pro',
        generation_config=genai.types.GenerationConfig(
            temperature=0.7,
            max_output_tokens=4096,
            response_mime_type="application/json"
        ),
        safety_settings={
            types.HarmCategory.HARM_CATEGORY_HARASSMENT: types.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
            types.HarmCategory.HARM_CATEGORY_HATE_SPEECH: types.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
            types.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: types.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
            types.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: types.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
        },
        transport="rest" # Use REST for memory efficiency
    )
    
    try:
        # --- 1. GENERATE TEXT ---
        # This call uses the API Key (transport="rest")
        response = model.generate_content(prompt)
        
        if not response.candidates or not response.candidates[0].content.parts:
            finish_reason = response.candidates[0].finish_reason.name if response.candidates else "UNKNOWN"
            logging.error(f"Model blocked text generation. Finish Reason: {finish_reason}")
            raise ValueError(f"Model blocked text generation. Finish Reason: {finish_reason}")
        
        response_text = response.text.strip()
        data = JSON_LOADER(response_text)
        
        insight = BusinessInsight(**data)

        # --- 2. GENERATE IMAGE ---
        # This call uses the Service Account (via generate_ai_image)
        
        image_filename = f"{timestamp}_insights.jpg" 
        
        image_success = generate_ai_image(
            prompt=insight.image_generation_prompt,
            image_filename=image_filename
        )

        if not image_success:
            print("   Falling back to placeholder image...")
            image_filename = None 
        
        return insight, hook_and_cta_text, image_filename
        
    except json.JSONDecodeError as e:
        logging.error(f"❌ JSON parsing error: {e}")
        try:
            logging.error(f"Response snippet: {response.text[:500]}")
        except:
            pass
        raise
    except Exception as e:
        logging.error(f"❌ Gemini API error: {e}")
        raise
# === END STEP 5 ===


# === (FIXED) STEP 6: SAVE OUTPUTS (FULL TEMPLATE) ===
def save_outputs(insight: BusinessInsight, hook_and_cta_text: str, image_filename: str | None):
    """
    Saves all output files (HTML, DOCX, XLSX).
    """
    date_str = datetime.now().strftime("%d %B %Y")
    
    # --- Define filenames ---
    html_filename_base = f"{timestamp}_insights"
    html_filename_full = f"{html_filename_base}.html"
    
    # --- Handle image filename ---
    if image_filename:
        # We already generated the image with the correct name (e.g., "..._insights.jpg")
        html_image_src = escape(image_filename) 
    else:
        # Create a placeholder if one doesn't exist
        placeholder_path = PUBLIC_HTML_DIR / "default.jpg"
        if not placeholder_path.exists():
            try:
                # Simple 16:9 grey SVG placeholder
                svg_placeholder = (
                    '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 16 9" width="1280" height="720">'
                    '<rect width="16" height="9" fill="#E5E7EB" />'
                    '<text x="50%" y="50%" dominant-baseline="middle" text-anchor="middle" '
                    'font-family="Inter, sans-serif" font-size="1" fill="#6B7280">Image Pending</text>'
                    '</svg>'
                )
                with open(placeholder_path, "w") as f:
                    f.write(svg_placeholder)
                print(f"✅ Created placeholder image: {placeholder_path}")
            except Exception as e:
                print(f"⚠️ Could not create placeholder image: {e}")
        
        # Use the relative path to the default image
        html_image_src = "default.jpg" 

    # --- NEW: FULL HTML PAGE TEMPLATE ---
    html = f"""<!DOCTYPE html>
<html lang="en-AU">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{escape(insight.title)}</title>
    
    <meta name="description" content="{escape(insight.tile_summary)}">
    <meta property="og:title" content="{escape(insight.title)}">
    <meta property="og:description" content="{escape(insight.tile_summary)}">
    <meta property="og:type" content="article">
    <meta property="og:image" content="{html_image_src}">

    <meta name="tile-summary" content="{escape(insight.tile_summary)}">
    <meta name="image-filename" content="{html_image_src}">
    <meta name="image-alt" content="{escape(insight.image_alt_text)}">

    <link rel="icon" href="../assets/favicon.ico" type="image/x-icon">
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;700&family=Roboto+Mono&display=swap" rel="stylesheet">
    
    <link rel="stylesheet" href="../style.css">

    <style>
        .main-article-image {{
            width: 100%;
            height: auto;
            max-height: 400px;
            object-fit: cover;
            border-radius: var(--border-radius);
            margin-bottom: 2rem;
            border: 1px solid #E5E7EB;
            background-color: #E5E7EB; /* Shows if image is broken */
        }}
        .blog-post-content p {{ margin-bottom: 1.5rem; }}
        .blog-post-content ul {{ list-style-type: disc; padding-left: 1.5rem; margin-bottom: 1.5rem; }}
        .blog-post-content li {{ margin-bottom: 0.5rem; }}
    </style>
</head>
<body>
    <a href="#main-content" class="skip-link">Skip to main content</a>
    
    <header class="site-header">
        <div class="container">
            <a href="../index.html" class="logo" aria-label="DigitalABCs Home">
                <img src="../assets/logo.png" alt="DigitalABCs Logo" class="logo-img">
            </a>
            <button class="nav-toggle" aria-label="Open navigation menu" aria-expanded="false" aria-controls="main-nav-menu">
                <span class="hamburger-line"></span>
                <span class="hamburger-line"></span>
                <span class="hamburger-line"></span>
            </button>
            <nav class="main-nav" id="main-nav-menu">
                <ul>
                    <li><a href="../index.html">Home</a></li>
                    <li><a href="../about.html">About</a></li>
                    <li><a href="../services.html">Services</a></li>
                    <li><a href="insights.php" class="active">Insights</a></li>
                    <li><a href="../contact.html">Contact</a></li>
                    <li><a href="https://services.digitalabcs.com.au" target="_blank">Client Login</a></li>
                </ul>
            </nav>
        </div>
    </header>

    <main id="main-content">
        <section class="hero-small">
            <div class="container">
                <h1>{escape(insight.title)}</h1>
                <p class="subtitle">Published: {date_str}</p>
            </div>
        </section>

        <section class="insights-grid section-padding">
            <div class="container" style="max-width: 800px;">
                <article class="blog-post-content">
                    <img src="{html_image_src}" alt="{escape(insight.image_alt_text)}" class="main-article-image">
                    
                    <h2>The Big Picture: AI, Automation, and Your Power</h2>
                    <p class="summary-text" style="white-space: pre-wrap;">{escape(insight.summary)}</p>
                    
                    <h2>Your Action Plan: Practical AI & Automation Takeaways</h2>
                    <ul>{''.join(f'<li><strong>Time-Sensitive Action:</strong> {escape(u)}</li>' for u in insight.key_updates)}</ul>
                    
                    <div class="cta-box" style="border-color: var(--color-green-cta);">
                        <h2>Ready to Take Back Control?</h2>
                        <p class="cta-text" style="color: var(--color-green-cta); font-weight: bold;">{escape(hook_and_cta_text)}</p>
                        <button class="cta" style="background-color: var(--color-green-cta); color: white; border: none; padding: 10px 20px; border-radius: 4px; cursor: pointer; font-weight: bold; margin-top: 10px;" onclick="window.open('../contact.html', '_self')">Start Simplifying Your Business</button>
                    </div>
                    
                    <h2 style="margin-top: 2rem;">Sources</h2>
                    <ul>{''.join(f'<li><a href="{escape(s)}" target="_blank" rel="noopener noreferrer">{escape(s)}</a></li>' for s in insight.sources)}</ul>
                </article>
            </div>
        </section>
    </main>

    <footer class="site-footer">
        <div class="container">
            <div class="footer-grid">
                <div class="footer-col">
                    <h4>DigitalABCs</h4>
                    <p>Empowering Australian businesses through practical technology. Built on resilience, for resilience.</p>
                </div>
                <div class="footer-col">
                    <h4>Navigate</h4>
                    <ul>
                        <li><a href="../about.html">About Us</a></li>
                        <li><a href="../services.html">Our Services</a></li>
                        <li><a href="insights.php">Insights</a></li>
                        <li><a href="../contact.html">Contact</a></li>
                    </ul>
                </div>
                <div class="footer-col">
                    <h4>Legal</h4>
                    <ul>
                        <li><a href="../privacy.html">Privacy Policy</a></li>
                        <li><a href="../terms.html">Terms of Service</a></li>
                    </ul>
                </div>
                <div class="footer-col">
                    <h4>Connect</h4>
                    <p>info@digitalabcs.com.au<br>Toongabbie, NSW, Australia</p>
                </div>
            </div>
            <div class="footer-bottom">
                <p>&copy; 2025 DigitalABCs. All rights reserved.</p>
            </div>
        </div>
    </footer>

    <script>
        document.addEventListener('DOMContentLoaded', function () {{
            const navToggle = document.querySelector('.nav-toggle');
            const mainNav = document.querySelector('.main-nav');
            if (navToggle && mainNav) {{
                navToggle.addEventListener('click', function () {{
                    const isActive = mainNav.classList.toggle('is-active');
                    navToggle.setAttribute('aria-expanded', isActive);
                    navToggle.setAttribute('aria-label', isActive ? 'Close navigation menu' : 'Open navigation menu');
                }});
            }}
        }});
    </script>
</body>
</html>"""
    
    # --- Define Archive/Public Filenames ---
    archive_html_file = ARCHIVE_OUTPUT_DIR / f"{html_filename_base}.html"
    public_html_file = PUBLIC_HTML_DIR / f"{html_filename_base}.html"
    archive_docx_file = ARCHIVE_OUTPUT_DIR / f"{html_filename_base}.docx"
    archive_excel_file = ARCHIVE_OUTPUT_DIR / f"{html_filename_base}.xlsx"

    
    # --- Save HTML file ---
    with open(archive_html_file, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"✅ HTML archived: {archive_html_file}")

    # Copy the HTML file to the PUBLIC web directory
    try:
        shutil.copyfile(archive_html_file, public_html_file)
        print(f"✅ HTML published: {public_html_file}")
    except Exception as e:
        print(f"❌ FAILED to copy HTML to public directory: {e}")

    # --- Save Word Doc ---
    doc = DocxDocument()
    doc.add_heading(insight.title, 0)
    doc.add_paragraph(f"Published: {date_str}")
    
    doc.add_heading("Tile Summary (for website grid)", level=2)
    doc.add_paragraph(insight.tile_summary)
    doc.add_heading("Image Generation Prompt", level=2)
    doc.add_paragraph(insight.image_generation_prompt)
    doc.add_heading("Image Alt Text (WCAG)", level=2)
    doc.add_paragraph(insight.image_alt_text)
    
    doc.add_heading("The Big Picture: AI, Automation, and Your Power", level=1)
    doc.add_paragraph(insight.summary)
    doc.add_heading("Your Action Plan: Practical AI & Automation Takeaways", level=1)
    for update in insight.key_updates:
        doc.add_paragraph(f"Time-Sensitive Action: {update}", style='List Bullet')
    doc.add_heading("Ready to Take Back Control?", level=1)
    doc.add_paragraph(hook_and_cta_text)
    doc.add_heading("Sources", level=1)
    for src in insight.sources:
        doc.add_paragraph(src, style='List Bullet')
    doc.save(archive_docx_file)
    print(f"✅ Word doc saved: {archive_docx_file}")

    # --- Save Excel Sheet ---
    df = pd.DataFrame({
        "AI/Automation Action Item": insight.key_updates,
        "Tile Summary": [insight.tile_summary] + [""] * (len(insight.key_updates) - 1),
        "Image Prompt": [insight.image_generation_prompt] + [""] * (len(insight.key_updates) - 1),
        "Alt Text": [insight.image_alt_text] + [""] * (len(insight.key_updates) - 1)
    })
    df.to_excel(archive_excel_file, index=False)
    print(f"✅ Excel saved: {archive_excel_file}")
# === END STEP 6 ===


# === STEP 7: ACCESSIBILITY AUDIT ===
# def audit_accessibility():
#     print("🔍 Running WCAG audit...")
#     options = webdriver.ChromeOptions()
#     options.add_argument("--headless")
#     options.add_argument("--no-sandbox")
#     options.add_argument("--disable-dev-shm-usage")
    
#     try:
#         # 🟢 FIX: Need to define ARCHIVE_HTML_FILE globally
#         # or pass it as an argument
#         global_archive_html_file = ARCHIVE_OUTPUT_DIR / f"{timestamp}_insights.html"
#         driver = webdriver.Chrome(options=options)
#         driver.get(f"file:///{global_archive_html_file.resolve()}")
#         axe = Axe(driver)
#         axe.inject()
#         results = axe.run(options={'runOnly': {'type': 'tag', 'values': ['wcag2a', 'wcag2aa', 'wcag21a', 'wcag21aa']}})
#         driver.quit()
        
#         if results["violations"]:
#             print("⚠️  WCAG issues found:")
#             for v in results["violations"]:
#                 print(f" - {v['description']}")
#                 print(f"  Affected nodes: {[n['html'] for n in v['nodes'][:2]]}...")
#         else:
#             print("✅ WCAG 2.1 AA compliant!")
#     except Exception as e:
#         print(f"⚠️  Accessibility Audit skipped (WebDriver issue in environment?): {e}")

# === (FIXED) MAIN ===
def main():
    print("🇦🇺 DigitalABCs Insights Agent (Gemini Native SDK)")
    
    # Setup
    kb = create_compliance_kb()
    news = fetch_recent_news()
    print(f"📰 Found {len(news)} articles in the last 7 days.")
    
    # Generate
    try:
        # 🟢 FIX: Unpack all 3 returned values
        insight, hook_text, image_filename = generate_insights(news, kb)
        print(f"🧠 Generated: {insight.title}")
    except Exception as e:
        print(f"❌ Generation failed: {e}")
        print(f"❌ Error type: {type(e).__name__}")
        import traceback
        print("❌ Full traceback:")
        traceback.print_exc()
        sys.exit(1)
    
    # Output
    # 🟢 FIX: Pass all 3 values to save_outputs
    save_outputs(insight, hook_text, image_filename)
    
    # 🟢 NOTE: Audit is commented out by default to save resources.
    # audit_accessibility() 
    
    print(f"\n🎉 Success! Outputs archived in: {ARCHIVE_OUTPUT_DIR.resolve()}")
    # 🟢 FIX: Corrected public path in print message
    print(f"🎉 Public HTML published to: {PUBLIC_HTML_DIR.resolve()}")


if __name__ == "__main__":
    main()