# insights_agent.py

# DigitalABCs Compliant AI Agent — Gemini API Version

# Empowerment | Compliance | Practicality | Privacy


# -----------------------------------------------------------------

# --- 🔴 START: ENVIRONMENT FIX REMOVED 🔴 -
# The cPanel-specific GRPC environment variables are REMOVED 

# as they restrict performance and are unnecessary on Windows 11.

# -----------------------------------------------------------------



# -----------------------------------------------------------------

# --- IMPORTS ---

# -----------------------------------------------------------------
import logging

import sys

from pathlib import Path 

import shutil # To handle local file movements

import base64 # To create the placeholder image on failure

from html import escape # To sanitize HTML output

import os # Keep os import at the top


# --- 🟢 NEW: UPLOAD LIBRARIES 🟢 -

import requests # For uploading the final files back to the server.

# ------------------------------------


# --- ▼▼▼ START: SERVICE ACCOUNT AUTHENTICATION ▼▼▼ ---

# This block MUST come before you import 'vertexai'

SCRIPT_DIR = Path(__file__).resolve().parent


# This should be the name of the JSON key file you uploaded

KEY_FILE_PATH = SCRIPT_DIR / "gen-lang-client-0944750405-a5911516ca25.json"


if not KEY_FILE_PATH.exists():

    print(f"❌ Error: Authentication key file not found at {KEY_FILE_PATH}")

    print("❌ Please ensure the service account has the 'Vertex AI User' role.")

    print("❌ Please download your Google Cloud service account JSON key,")

    print("❌ upload it to this directory, and ensure the filename matches.")

    sys.exit(1)


os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = str(KEY_FILE_PATH)

# --- ▲▲▲ END: SERVICE ACCOUNT AUTHENTICATION ▲▲▲ ---



import json

from datetime import datetime, timedelta

from typing import List, Annotated


# ADDED: Vertex AI Imports for Image Generation

import vertexai

try:

    from vertexai.preview.vision_models import ImageGenerationModel 

except ImportError:

    print("WARNING: vertexai library not fully found. Image generation might fail.")

ImageGenerationModel = None


# FIX: Attempt to import orjson for robust JSON decoding

try:

    import orjson  # optional fast JSON parser

    JSON_LOADER = orjson.loads

except Exception:

    JSON_LOADER = json.loads



# Load environment variables

from dotenv import load_dotenv

load_dotenv()


# Structured output for trust & simplicity

from pydantic import BaseModel, Field


# Gemini SDK

import google.generativeai as genai  # <--- FIX 1

from google.generativeai import types # <--- FIX 2


# Pure Chroma for private RAG (no LangChain)


# 🟢 FIX: Set the default ChromaDB configuration class to ensure local mode.
# This resolves the 12 validation errors by ignoring network configuration checks.

os.environ["CHROMA_SETTINGS_CLASS"] = "chromadb.config.Settings" 


import chromadb

from chromadb.utils import embedding_functions


# Document & news handling
from dateutil import parser as date_parser
import feedparser

from docx import Document as DocxDocument


# Practical outputs
import pandas as pd


# Ensure clean environment

os.environ["TOKENIZERS_PARALLELISM"] = "false"

# ... (Configuration, functions, and main logic remain the same) ...


# === (FIXED) STEP 6: SAVE OUTPUTS (FULL TEMPLATE) ===

def save_outputs(insight: BusinessInsight, hook_and_cta_text: str, image_filename: str | None):
    """
    Saves all output files (HTML, DOCX, XLSX) locally, and prepares public files.
    Returns paths to public files for upload.
    """

    date_str = datetime.now().strftime("%d %B %Y")

    # --- Define filenames ---

    html_filename_base = f"{timestamp}_insights"

    html_filename_full = f"{html_filename_base}.html"

    metadata_filename_full = f"{html_filename_base}.json"


    # --- Handle image filename ---

    if image_filename:

        # The image is already saved to PUBLIC_HTML_DIR (staging)

        html_image_src = escape(image_filename)

        public_image_file = PUBLIC_HTML_DIR / image_filename

    else:

        # Create a placeholder in the local staging folder if one doesn't exist

        placeholder_path = PUBLIC_HTML_DIR / "default.jpg"
            

        if not placeholder_path.exists():

            try:

                # Save the simple SVG placeholder to the staging directory

                svg_placeholder = (

                    '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 16 9" width="1280" height="720">'

                    '<rect width="16" height="9" fill="#E5E7EB" />'

                    '<text x="50%" y="50%" dominant-baseline="middle" text-anchor="middle" '

                    'font-family="Inter, sans-serif" font-size="1" fill="#6B7280">Image Pending</text>'

                    '</svg>'

                )

                with open(placeholder_path, "w", encoding="utf-8") as f:

                    f.write(svg_placeholder)

                    print(f"✅ Created placeholder image: {placeholder_path}")

            except Exception as e:

                print(f"⚠️ Could not create placeholder image: {e}")


        # Use the relative path to the default image

        html_image_src = "default.jpg" 

        public_image_file = placeholder_path # Pass the path to the placeholder file


    # --- FIX 3: Re-insert the full HTML template below ---

    html = f"""<!DOCTYPE html>

<html lang="en-AU">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{escape(insight.title)}</title>

        <meta name="description" content="{escape(insight.tile_summary)}">
        <meta property="og:title" content="{escape(insight.title)}">
        <meta property="og:description" content="{escape(insight.tile_summary)}">
        <meta property="og:type" content="article">
        <meta property="og:image" content="{html_image_src}">

        <meta name="tile-summary" content="{escape(insight.tile_summary)}">
        <meta name="image-filename" content="{html_image_src}">
        <meta name="image-alt" content="{escape(insight.image_alt_text)}">

        <link rel="icon" href="../assets/favicon.ico" type="image/x-icon">
        <link rel="preconnect" href="https://fonts.googleapis.com">
        <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
        <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;700&family=Roboto+Mono&display=swap" rel="stylesheet">


        <link rel="stylesheet" href="../style.css">

        <style>
            .main-article-image {{
                width: 100%;
                height: auto;
                max-height: 400px;
                object-fit: cover;
                border-radius: var(--border-radius);
                margin-bottom: 2rem;
                border: 1px solid #E5E7EB;
                background-color: #E5E7EB; /* Shows if image is broken */
            }}
            .blog-post-content p {{ margin-bottom: 1.5rem; }}
            .blog-post-content ul {{ list-style-type: disc; padding-left: 1.5rem; margin-bottom: 1.5rem; }}
            .blog-post-content li {{ margin-bottom: 0.5rem; }}
        </style>
    </head>
    <body>
        <a href="#main-content" class="skip-link">Skip to main content</a>
        <header class="site-header">
            <div class="container">
                <a href="../index.html" class="logo" aria-label="DigitalABCs Home">
                    <img src="../assets/logo.png" alt="DigitalABCs Logo" class="logo-img">
                </a>
                <button class="nav-toggle" aria-label="Open navigation menu" aria-expanded="false" aria-controls="main-nav-menu">
                    <span class="hamburger-line"></span>
                    <span class="hamburger-line"></span>
                    <span class="hamburger-line"></span>
                </button>
                <nav class="main-nav" id="main-nav-menu">
                    <ul>
                        <li><a href="../index.html">Home</a></li>
                        <li><a href="../about.html">About</a></li>
                        <li><a href="../services.html">Services</a></li>
                        <li><a href="insights.php" class="active">Insights</a></li>
                        <li><a href="../contact.html">Contact</a></li>
                        <li><a href="https://services.digitalabcs.com.au" target="_blank">Client Login</a></li>
                    </ul>
                </nav>
            </div>
        </header>

        <main id="main-content">
            <section class="hero-small">
                <div class="container">
                    <h1>{escape(insight.title)}</h1>
                    <p class="subtitle">Published: {date_str}</p>
                </div>
            </section>

            <section class="insights-grid section-padding">
                <div class="container" style="max-width: 800px;">
                    <article class="blog-post-content">
                        <img src="{html_image_src}" alt="{escape(insight.image_alt_text)}" class="main-article-image">
                        
                        <h2>The Big Picture: AI, Automation, and Your Power</h2>

                        <p class="summary-text" style="white-space: pre-wrap;">{escape(insight.summary)}</p>
                    
                        <h2>Your Action Plan: Practical AI & Automation Takeaways</h2>
                        <ul>{''.join(f'<li><strong>Time-Sensitive Action:</strong> {escape(u)}</li>' for u in insight.key_updates)}</ul>
                        <div class="cta-box" style="border-color: var(--color-green-cta);">
                            <h2>Ready to Take Back Control?</h2>
                            <p class="cta-text" style="color: var(--color-green-cta); font-weight: bold;">{escape(hook_and_cta_text)}</p>
                            <button class="cta" style="background-color: var(--color-green-cta); color: white; border: none; padding: 10px 20px; border-radius: 4px; cursor: pointer; font-weight: bold; margin-top: 10px;" onclick="window.open('../contact.html', '_self')">Start Simplifying Your Business</button>
                        </div>
                
                        <h2 style="margin-top: 2rem;">Sources</h2>
                        <ul>{''.join(f'<li><a href="{escape(s)}" target="_blank" rel="noopener noreferrer">{escape(s)}</a></li>' for s in insight.sources)}</ul>
                    </article>
                </div>
            </section>
        </main>

        <footer class="site-footer">
            <div class="container">
                <div class="footer-grid">
                    <div class="footer-col">
                        <h4>DigitalABCs</h4>
                        <p>Empowering Australian businesses through practical technology. Built on resilience, for resilience.</p>
                    </div>
                    <div class="footer-col">
                        <h4>Navigate</h4>
                        <ul>
                            <li><a href="../about.html">About Us</a></li>
                            <li><a href="../services.html">Our Services</a></li>
                            <li><a href="insights.php">Insights</a></li>
                            <li><a href="../contact.html">Contact</a></li>
                            <li><a href="https://services.digitalabcs.com.au" target="_blank">Client Login</a></li>
                        </ul>
                    </div>
                    <div class="footer-col">
                        <h4>Legal</h4>
                        <ul>
                            <li><a href="../privacy.html">Privacy Policy</a></li>
                            <li><a href="../terms.html">Terms of Service</a></li>
                        </ul>
                    </div>
                    <div class="footer-col">
                        <h4>Connect</h4>
                        <p>info@digitalabcs.com.au<br>Toongabbie, NSW, Australia</p>
                    </div>
                    <div class="footer-bottom">
                        <p>&copy; 2025 DigitalABCs. All rights reserved.</p>
                    </div>
                </div>
            </div>
        </footer>

        <script>
            document.addEventListener('DOMContentLoaded', function () {{
                const navToggle = document.querySelector('.nav-toggle');
                const mainNav = document.querySelector('.main-nav');
                if (navToggle && mainNav) {{
                    navToggle.addEventListener('click', function () {{
                        const isActive = mainNav.classList.toggle('is-active');
                        navToggle.setAttribute('aria-expanded', isActive);
                        navToggle.setAttribute('aria-label', isActive ? 'Close navigation menu' : 'Open navigation menu');
                    }});
                }}
            }});
        </script>
    </body>
    </html>"""


    # --- Define Archive/Public Filenames ---

    archive_html_file = ARCHIVE_OUTPUT_DIR / html_filename_full

    public_html_file = PUBLIC_HTML_DIR / html_filename_full

    archive_docx_file = ARCHIVE_OUTPUT_DIR / f"{html_filename_base}.docx"

    archive_excel_file = ARCHIVE_OUTPUT_DIR / f"{html_filename_base}.xlsx"

    public_metadata_file = PUBLIC_HTML_DIR / metadata_filename_full # New metadata file



    # --- Save HTML file (Archive & Public Staging) ---

    with open(archive_html_file, "w", encoding="utf-8") as f:

        f.write(html)

    print(f"✅ HTML archived: {archive_html_file}")


    # 🟢 UPDATE: Save HTML to public staging (no shutil.copyfile needed

    with open(public_html_file, "w", encoding="utf-8") as f:

        f.write(html)

    print(f"✅ HTML saved to local staging: {public_html_file}")



    # 🟢 NEW: Save JSON Metadata file (Crucial for website to index the new article

    metadata = {

        "title": insight.title,

        "tile_summary": insight.tile_summary,

        "image_filename": html_image_src,

        "image_alt_text": insight.image_alt_text,

        "link_path": html_filename_full,

        "timestamp": timestamp,

    }

    with open(public_metadata_file, "w", encoding="utf-8") as f:

        json.dump(metadata, f, indent=4)

    print(f"✅ Metadata JSON saved to local staging: {public_metadata_file}")



    # --- Save Word Doc ---

    doc = DocxDocument()

    doc.add_heading(insight.title, 0)

    doc.add_paragraph(f"Published: {date_str}")


    doc.add_heading("Tile Summary (for website grid)", level=2)

    doc.add_paragraph(insight.tile_summary)

    doc.add_heading("Image Generation Prompt", level=2)

    doc.add_paragraph(insight.image_generation_prompt)

    doc.add_heading("Image Alt Text (WCAG)", level=2)

    doc.add_paragraph(insight.image_alt_text)


    doc.add_heading("The Big Picture: AI, Automation, and Your Power", level=1)

    doc.add_paragraph(insight.summary)

    doc.add_heading("Your Action Plan: Practical AI & Automation Takeaways", level=1)

    for update in insight.key_updates:

        doc.add_paragraph(f"Time-Sensitive Action: {update}", style='List Bullet')

    doc.add_heading("Ready to Take Back Control?", level=1)

    doc.add_paragraph(hook_and_cta_text)

    doc.add_heading("Sources", level=1)

    for src in insight.sources:

        doc.add_paragraph(src, style='List Bullet')

    doc.save(archive_docx_file)

    print(f"✅ Word doc saved: {archive_docx_file}")


    # --- Save Excel Sheet ---

    df = pd.DataFrame({

        "AI/Automation Action Item": insight.key_updates,

        "Tile Summary": [insight.tile_summary] + [""] * (len(insight.key_updates) - 1),

        "Image Prompt": [insight.image_generation_prompt] + [""] * (len(insight.key_updates) - 1),

        "Alt Text": [insight.image_alt_text] + [""] * (len(insight.key_updates) - 1)

    })

    df.to_excel(archive_excel_file, index=False)

    print(f"✅ Excel saved: {archive_excel_file}")

# === END STEP 6 ===



# === (FIXED) MAIN ===

def main():
    print("🇦🇺 DigitalABCs Insights Agent (Gemini Native SDK)")


# Setup

kb = create_compliance_kb()

news = fetch_recent_news()

print(f"📰 Found {len(news)} articles in the last 7 days.")


# Generate

try:

    insight, hook_text, image_filename = generate_insights(news, kb)

    print(f"🧠 Generated: {insight.title}"
)
except Exception as e:

    print(f"❌ Generation failed: {e}")

    print(f"❌ Error type: {type(e).__name__}")

# import traceback # Commented out for cleaner output

# traceback.print_exc()

sys.exit(1)


# Output and Local Save

html_path, image_path, metadata_path = save_outputs(insight, hook_text, image_filename)


# 🟢 NEW: Remote Upload Ste
p
print("\n--- Initiating Remote Upload ---")

upload_success = upload_files_to_server(html_path, image_path, metadata_path)

if upload_success:

    print("🎉 Remote deployment initiated successfully via API."
)
else:

    print("⚠️ Remote deployment failed. Files are saved locally for manual upload.")


print(f"\n🎉 Success! Outputs archived in: {ARCHIVE_OUTPUT_DIR.resolve()}"
)
print(f"🎉 Public files staged at: {PUBLIC_HTML_DIR.resolve()}")



if __name__ == "__main__":

    main()